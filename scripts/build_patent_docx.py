#!/usr/bin/env python3
"""以 1110 改稿为底稿复制并打补丁，生成新专利 docx。"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from utils.patent_experiment_tables import paired_table_rows  # noqa: E402

ORIGINAL_1110 = ROOT / "一种基于强化学习的多智能体协同优化方法及系统（发明）1110改稿1.docx"
NEW_PATH = ROOT / "一种面向分布式黑盒优化的强化学习协同惩罚优化系统及方法.docx"

NEW_TITLE = "一种面向分布式黑盒优化的强化学习协同惩罚优化系统及方法"

NEW_ABSTRACT = (
    "本发明公开了一种面向分布式黑盒优化的强化学习协同惩罚优化系统及方法，该方法包括："
    "根据MACPO分布式优化框架定义智能体的邻居集合与共享变量索引集，计算冲突强度和冲突趋势；"
    "进行门控决策，并构建通信必要条件，满足则触发通信协商；对共享变量索引集进行重要性筛选，"
    "根据筛选结果进行通信协商，并通过策略网络根据冲突强度和冲突趋势输出动作；"
    "根据策略网络输出的动作更新惩罚权重并结合目标动作，更新网络参数，实现多智能体协同优化。"
    "配对仿真实验表明：在多区域阀点经济调度（13机、1800MW）场景通信触发率可由100%降至约11.3%"
    "（降幅88.7%）且最优纯目标相对变化约0.01%；"
    "在资源约束调度场景通信降幅约80%且纯目标有所改善。"
    "本发明能够实现多智能体协同优化，提高数值稳定性，并通过强化学习机制实现动态环境下的自适应优化，"
    "可应用于区域互联电力经济调度、多区域阀点经济调度及资源约束型分布式协同优化。"
)

PAIRING_SECTION = [
    "实施例五、工程场景配对验证实验",
    "在相同随机种子、相同函数评估预算上限及相同MPI进程数条件下，将MACPO基线与本发明全功能配置"
    "于多区域阀点经济调度（13机、1800MW）、资源约束型分布式调度、"
    "电动汽车充放电协同调度场景各重复运行10次。通信触发率定义为触发跨节点协商的外循环轮次占比；"
    "最优纯目标为全程最优且不含惩罚项的全局目标均值。配对协议保证双方评估次数上限一致。",
    "表1为上述场景的配对实验结果：",
]

TABLE1_ANALYSIS = (
    "表1表明：多区域阀点经济调度（13机、1800MW）在通信降幅88.7%时最优纯目标相对MACPO变化约0.01%；"
    "资源约束调度与电动汽车充放电调度在通信降幅80.0%～87.3%时纯目标改善约1.6%～1.8%。"
    "阀点经济调度外循环轮次较多，fail-safe机制累计触发协商，配合跳过通信时的局部最优合并与共享变量同步，"
    "可在降低通信的同时维持边界一致性；资源约束场景在约束紧张阶段冲突指数升高，门控触发更频繁，"
    "因而更易保持或改善纯目标。墙钟耗时可能因策略网络在线推理而高于基线，但核心评价指标为"
    "同评估预算下的通信触发率与纯目标质量。"
)


def _set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "000000")
        borders.append(elem)
    tbl_pr.append(borders)


def _write_cell(cell, text: str) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _make_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    _set_table_borders(table)
    for i, h in enumerate(headers):
        _write_cell(table.rows[0].cells[i], h)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            _write_cell(cells[i], val)
    return table


def _find_para(doc: Document, needle: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    return None


def _replace_in_doc(doc: Document, old: str, new: str) -> int:
    n = 0
    for p in doc.paragraphs:
        if old in p.text:
            p.text = p.text.replace(old, new)
            n += 1
    return n


def _insert_paragraph_after(anchor: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._element.addnext(new_p)
    para = Paragraph(new_p, anchor._parent)
    para.add_run(text)
    return para


def _insert_paragraph_before(anchor: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._element.addprevious(new_p)
    para = Paragraph(new_p, anchor._parent)
    para.add_run(text)
    return para


def _insert_table_before(doc: Document, anchor: Paragraph, headers: list[str], rows: list[list[str]]) -> None:
    table = _make_table(doc, headers, rows)
    tbl_el = table._tbl
    doc.element.body.remove(tbl_el)
    anchor._element.addprevious(tbl_el)


def _insert_blocks_before(doc: Document, anchor: Paragraph, blocks: list) -> None:
    """blocks 自上而下顺序与列表一致，整体插在 anchor 之前。"""
    insert_point = anchor._element
    for block in reversed(blocks):
        if isinstance(block, str):
            new_p = OxmlElement("w:p")
            insert_point.addprevious(new_p)
            para = Paragraph(new_p, anchor._parent)
            para.add_run(block)
            insert_point = new_p
        elif block[0] == "table":
            _, headers, rows = block
            table = _make_table(doc, headers, rows)
            tbl_el = table._tbl
            doc.element.body.remove(tbl_el)
            insert_point.addprevious(tbl_el)
            insert_point = tbl_el


def _append_claim_before_system(doc: Document) -> None:
    anchor = _find_para(doc, "一种基于强化学习的多智能体协同优化系统，其特征在于，包括以下模块：")
    if anchor is None:
        return
    claims = [
        "根据权利要求1所述一种基于强化学习的多智能体协同优化方法，其特征在于，在未触发通信协商的外循环迭代中，"
        "执行局部最优合并与重叠共享变量同步，以维持共享边界一致性。",
        "根据权利要求1所述一种基于强化学习的多智能体协同优化方法，其特征在于，所述门控决策配置fail-safe机制："
        "当连续k轮未触发通信协商时强制触发一次协商，k为预设正整数。",
        "根据权利要求1所述一种基于强化学习的多智能体协同优化方法，其特征在于，所述方法应用于区域互联电力经济调度、"
        "多区域阀点经济调度或资源约束型分布式协同优化中的至少一种场景。",
    ]
    for text in reversed(claims):
        _insert_paragraph_before(anchor, text)


def build_new_docx() -> Document:
    if not ORIGINAL_1110.exists():
        raise FileNotFoundError(f"找不到原始专利底稿：{ORIGINAL_1110}")

    doc = Document(str(ORIGINAL_1110))

    # 摘要、标题
    if doc.paragraphs:
        doc.paragraphs[0].text = NEW_ABSTRACT
    title_p = _find_para(doc, "一种基于强化学习的多智能体协同优化方法及系统")
    if title_p is not None:
        title_p.text = NEW_TITLE

    # 技术领域
    _replace_in_doc(
        doc,
        "本发明涉及多智能体强化学习技术领域，尤其涉及一种基于强化学习的多智能体协同优化方法及系统。",
        "本发明涉及多智能体强化学习与分布式优化技术领域，尤其涉及一种面向分布式黑盒优化的"
        "强化学习协同惩罚优化系统及方法，可应用于电力系统经济调度等资源约束型协同优化场景。",
    )

    # 背景技术补充
    bg = _find_para(doc, "这些问题限制了多智能体强化学习在异构、动态环境中的应用效率。")
    if bg is not None:
        bg.text = (
            "这些问题限制了多智能体协同优化在异构、动态环境中的应用效率。"
            "工程实践中，区域互联电力调度、多区域阀点经济调度及资源约束型分布式协同优化等场景"
            "普遍存在目标函数不可解析、评估预算受限、共享边界变量须保持一致等约束。"
            "多智能体强化学习需要智能体间高效协作，同时学习最优策略，"
            "但相关方法往往忽略了个体化自适应和按需通信的需求。"
        )

    # 发明内容首段
    _replace_in_doc(
        doc,
        "为了解决上述技术问题，本发明的目的是提供一种基于强化学习的多智能体协同优化方法及系统，"
        "能够实现多智能体协同优化，提高数值稳定性，并通过强化学习机制实现了动态环境下的自适应优化。",
        "为了解决上述技术问题，本发明的目的是提供一种面向分布式黑盒优化的强化学习协同惩罚优化系统及方法，"
        "能够实现多智能体协同优化，在降低跨节点通信触发率的同时保持相当的解质量，"
        "提高数值稳定性，并通过强化学习机制实现动态环境下的自适应优化。",
    )

    # 有益效果补充定量表述
    effect = _find_para(doc, "本发明方法及系统的有益效果是：")
    if effect is not None:
        effect.text += (
            "配对实验进一步表明：多区域阀点经济调度(13机、1800MW)场景通信触发率由100%降至11.3%（降幅88.7%），"
            "最优纯目标相对变化约0.01%；资源约束型分布式调度场景通信降幅约80.0%，最优纯目标改善约1.8%。"
        )

    apply_shared_text_patches(doc)

    # 权利要求补充
    _append_claim_before_system(doc)

    # 配对实验表1（插在结语段之前）
    anchor = _find_para(doc, "以上是对本发明的较佳实施进行了具体说明")
    if anchor is not None and _find_para(doc, "表1为上述场景的配对实验结果：") is None:
        headers, rows = paired_table_rows()
        blocks: list = list(PAIRING_SECTION)
        if headers and rows:
            blocks.append(("table", headers, rows))
        blocks.append(TABLE1_ANALYSIS)
        _insert_blocks_before(doc, anchor, blocks)

    return doc


FIGURE_CAPTION_1 = "图1是本发明协同优化方法的步骤示意图（S100～S400）；"
FIGURE_CAPTION_2 = "图2是本发明多智能体协同优化系统的结构框图（模块201～204）；"
FIGURE_CAPTION_3 = (
    "图3是本发明具体实施例提供的外循环协同优化算法流程图"
    "（含跳过通信时的边界同步、强化学习更新及与图5、图6的衔接）；"
)
FIGURE_CAPTION_4 = "图4是本发明具体实施例提供的策略网络的结构示意图；"
FIGURE_CAPTION_5 = "图5是本发明具体实施例提供的门控通信判定流程图（含fail-safe强制触发）；"
FIGURE_CAPTION_6 = "图6是本发明具体实施例提供的智能变量筛选机制示意图。"

FIG3_BODY_OLD = (
    "因此，如图3所示，本发明实施例首先初始化多个智能体的局部种群、策略网络和变量重要性分数；"
    "进而每个智能体独立进行群体演化，获得局部最优解；进一步计算冲突强度和冲突趋势，其中, ；"
    "基于冲突指数进行门控决策，若且满足最小间隔和阶段频率抽样，则触发通信协商；"
    "判断是否触发通信门控，若满足最小间隔、自适应阈值和阶段频率抽样，则进入协商流程；"
    "在协商流程中，对共享变量进行Top-R%重要性筛选，仅让筛选后的变量参与逐维试探和单侧扰动冲突检测；"
    "进一步根据协商结果更新惩罚开关和共识解；每个智能体使用策略网络根据状态向量输出动作，更新惩罚权重；"
    "最后根据奖励信号进行策略网络的在线强化学习，包括经验回放和优先级采样。"
)

FIG3_BODY_NEW = (
    "因此，如图3所示，本发明外循环协同优化算法包括：初始化多个智能体的局部种群、策略网络和变量重要性分数；"
    "各智能体独立进行群体演化以获得局部最优解；计算冲突强度和冲突趋势；"
    "进行门控决策（判定逻辑如图5所示）；若触发通信协商，则先对共享变量进行Top-R%重要性筛选"
    "（如图6所示），仅让筛选后的变量参与逐维试探和单侧扰动冲突检测，并更新共识解与惩罚开关；"
    "若未触发通信协商，则执行局部最优合并与重叠共享变量同步以维持共享边界一致性；"
    "各智能体使用策略网络（如图4所示）根据包含冲突强度和冲突趋势的状态向量输出动作，更新惩罚权重；"
    "根据纯目标改善与冲突变化计算奖励信号，经经验回放与优先级采样更新策略网络参数；"
    "重复上述过程直至满足停止条件。"
)

SKIP_SYNC_NOTE = (
    "需要说明的是，当门控决策未触发通信协商时，各智能体执行局部最优合并与重叠共享变量同步，"
    "以维持共享边界一致性，具体流程如图3所示。"
)

POLICY_NET_OLD = (
    "如图4所示，构建单层前馈网络SimpleNet，输入层接收状态向量，线性层计算，"
    "其中w和b为网络参数，激活函数输出a = tanh(z)，限制动作范围在[-1,1]，"
    "网络参数通过反向传播更新，使用梯度下降法最小化损失。"
)

POLICY_NET_NEW = (
    "如图4所示，构建单层前馈网络SimpleNet，输入层接收状态向量"
    "（包括冲突强度CI与冲突趋势ΔCI），线性层计算z = w^T s + b，"
    "其中w和b为网络参数，激活函数输出a = tanh(z)，限制动作范围在[-1,1]，"
    "网络参数通过反向传播更新，使用梯度下降法最小化损失。"
)

REWARD_OLD = "奖励信号基于全局适应度定义，当时，否则，用于强化学习反馈。"

REWARD_NEW = (
    "奖励信号根据纯目标改善与冲突变化率计算，经tanh函数归一化至(-1,1)区间，用于强化学习反馈。"
)


def _replace_figure_captions(doc: Document) -> None:
    caption_map = {
        "图1是本发明": FIGURE_CAPTION_1,
        "图2是本发明": FIGURE_CAPTION_2,
        "图3是本发明": FIGURE_CAPTION_3,
        "图4是本发明": FIGURE_CAPTION_4,
        "图5是本发明": FIGURE_CAPTION_5,
        "图6是本发明": FIGURE_CAPTION_6,
    }
    for p in doc.paragraphs:
        text = p.text.strip()
        for prefix, new_caption in caption_map.items():
            if text.startswith(prefix) and text != new_caption:
                p.text = new_caption
                break


def _fix_pairing_section_order(doc: Document) -> None:
    """将误排的配对实验段落恢复为：实施例五 → 说明 → 表1引导 → 表格 → 分析。"""
    anchor_intro = _find_para(doc, "在相同随机种子、相同函数评估预算上限及相同MPI进程数条件下")
    anchor_table_line = _find_para(doc, "表1为上述场景的配对实验结果：")
    anchor_analysis = _find_para(doc, "表1表明：多区域阀点经济调度")
    anchor_closing = _find_para(doc, "以上是对本发明的较佳实施进行了具体说明")
    if not all([anchor_intro, anchor_table_line, anchor_analysis, anchor_closing]):
        return
    if _find_para(doc, "实施例五、工程场景配对验证实验") is not None:
        return

    # 当前错误顺序：分析 → 表 → 表1引导 → 说明 → 结语
    title = "实施例五、工程场景配对验证实验"
    intro_text = anchor_intro.text
    table_line = anchor_table_line.text
    analysis_text = anchor_analysis.text

    tbl_el = None
    for child in doc.element.body:
        if child.tag.endswith("tbl"):
            tbl_el = child
            break

    anchor_analysis._element.getparent().remove(anchor_analysis._element)
    if tbl_el is not None:
        tbl_el.getparent().remove(tbl_el)
    anchor_table_line._element.getparent().remove(anchor_table_line._element)
    anchor_intro._element.getparent().remove(anchor_intro._element)

    blocks: list = [title, intro_text, table_line]
    if tbl_el is not None:
        blocks.append(("table_xml", tbl_el))
    blocks.append(analysis_text)

    insert_point = anchor_closing._element
    for block in reversed(blocks):
        if isinstance(block, str):
            new_p = OxmlElement("w:p")
            insert_point.addprevious(new_p)
            para = Paragraph(new_p, anchor_closing._parent)
            para.add_run(block)
            insert_point = new_p
        elif isinstance(block, tuple) and block[0] == "table_xml":
            insert_point.addprevious(block[1])
            insert_point = block[1]


def apply_shared_text_patches(doc: Document) -> None:
    """附图说明、具体实施方式等与附图对应的文字修订（可重复执行）。"""
    _replace_figure_captions(doc)

    # 门控 fail-safe（具体实施方式 S200 说明）
    if _find_para(doc, "如图5所示，通信门控包括三重条件：") is None:
        _replace_in_doc(
            doc,
            "需要说明的是，如图5所示，通信门控包括三重条件：",
            "需要说明的是，当连续k轮未触发通信协商时，fail-safe机制强制触发一次协商，以避免共享边界长期失步。"
            "如图5所示，通信门控包括三重条件：",
        )

    # 系统模块描述
    _replace_in_doc(
        doc,
        "第二模块202，用于根据智能体的邻居集合与共享变量索引集计算冲突指数并进行门控决策，"
        "并构建通信必要条件，满足则触发通信协商；",
        "第二模块202，用于根据智能体的邻居集合与共享变量索引集计算冲突指数并进行门控决策，"
        "并构建通信必要条件，满足则触发通信协商，并在连续k轮未触发通信时由fail-safe机制强制触发协商；",
    )
    _replace_in_doc(
        doc,
        "第二模块，用于根据智能体的邻居集合与共享变量索引集计算冲突指数并进行门控决策，"
        "并构建通信必要条件，满足则触发通信协商；",
        "第二模块，用于根据智能体的邻居集合与共享变量索引集计算冲突指数并进行门控决策，"
        "并构建通信必要条件，满足则触发通信协商，并在连续k轮未触发通信时由fail-safe机制强制触发协商；",
    )

    # 跳过通信时的边界同步（S200）
    if _find_para(doc, SKIP_SYNC_NOTE) is None:
        anchor = _find_para(doc, "如图5所示，通信门控包括三重条件：")
        if anchor is not None:
            _insert_paragraph_after(anchor, SKIP_SYNC_NOTE)

    # 图3 总述段
    if _find_para(doc, FIG3_BODY_OLD):
        _replace_in_doc(doc, FIG3_BODY_OLD, FIG3_BODY_NEW)
    elif _find_para(doc, "因此，如图3所示，本发明外循环协同优化算法包括：") is None:
        anchor = _find_para(doc, "因此，如图3所示，")
        if anchor is not None:
            anchor.text = FIG3_BODY_NEW

    # 图4 策略网络
    _replace_in_doc(doc, POLICY_NET_OLD, POLICY_NET_NEW)

    # 奖励信号
    _replace_in_doc(doc, REWARD_OLD, REWARD_NEW)

    _fix_pairing_section_order(doc)

    _patch_patent_experiment_text(doc)
    _sync_table1(doc)


def _patch_patent_experiment_text(doc: Document) -> None:
    """表1 场景、摘要与有益效果中的配对实验表述（去除 MAED2，明确 0.01%）。"""
    if doc.paragraphs and "本发明公开了一种面向分布式" in doc.paragraphs[0].text:
        doc.paragraphs[0].text = NEW_ABSTRACT

    _replace_in_doc(doc, "、多区域阀点经济调度（2机）", "")
    _replace_in_doc(doc, "、多区域阀点经济调度(2机)", "")

    pairing_intro = PAIRING_SECTION[1]
    for p in doc.paragraphs:
        if "将MACPO基线与本发明全功能配置" in p.text and "各重复运行10次" in p.text:
            if p.text != pairing_intro:
                p.text = pairing_intro
            break

    table_old_variants = [
        (
            "表1表明：多区域阀点经济调度在通信降幅85.8%～88.7%时纯目标与MACPO基本持平（≈0%）；"
            "资源约束调度与电动汽车充放电调度在通信降幅80.0%～87.3%时纯目标改善约1.6%～1.8%。"
        ),
        (
            "表1表明：多区域阀点经济调度在通信降幅85.8%～88.7%时目标值与MACPO基本持平（≈0%）；"
            "资源约束调度与电动汽车充放电调度在通信降幅80.0%～87.3%时目标值改善约1.6%～1.8%。"
        ),
        (
            "表1表明：多区域阀点经济调度（13机、1800MW）在通信降幅88.7%时最优纯目标相对MACPO变化约0.01%；"
            "资源约束调度与电动汽车充放电调度在通信降幅80.0%～87.3%时纯目标改善约1.6%～1.8%。"
        ),
    ]
    for old in table_old_variants:
        if old != TABLE1_ANALYSIS:
            _replace_in_doc(doc, old, TABLE1_ANALYSIS)
    for p in doc.paragraphs:
        if p.text.startswith("表1表明：") and ("85.8%～88.7%" in p.text or "≈0%" in p.text):
            p.text = TABLE1_ANALYSIS
            break

    effect_old_variants = [
        (
            "配对实验进一步表明：多区域阀点经济调度(13机)场景通信触发率由100%降至11.3%（降幅88.7%），"
            "最优纯目标变化≈0%；资源约束型分布式调度场景通信降幅约80.0%，最优纯目标改善约1.8%。"
        ),
        (
            "配对实验进一步表明：多区域阀点经济调度(13机、1800MW)场景通信触发率由100%降至11.3%（降幅88.7%），"
            "最优纯目标变化≈0%；资源约束型分布式调度场景通信降幅约80.0%，最优纯目标改善约1.8%。"
        ),
        (
            "配对实验进一步表明：多区域阀点经济调度(13机)场景通信触发率由完全触发降至11.3%（降幅88.7%），"
            "最优目标值变化基本保持一致；资源约束型分布式调度场景通信降幅约80.0%，最优纯目标改善约1.8%。"
        ),
    ]
    effect_new = (
        "配对实验进一步表明：多区域阀点经济调度(13机、1800MW)场景通信触发率由100%降至11.3%（降幅88.7%），"
        "最优纯目标相对变化约0.01%；资源约束型分布式调度场景通信降幅约80.0%，最优纯目标改善约1.8%。"
    )
    for old in effect_old_variants:
        _replace_in_doc(doc, old, effect_new)

    abstract_old = [
        "配对仿真实验表明：在多区域阀点经济调度场景通信触发率可由100%降至约11%～14%且最优纯目标变化≈0%；",
        "配对仿真实验表明：在多区域阀点经济调度场景通信触发率可由100%降至约11.3%（降幅88.7%）且最优纯目标变化≈0%；",
    ]
    abstract_new = (
        "配对仿真实验表明：在多区域阀点经济调度（13机、1800MW）场景通信触发率可由100%降至约11.3%"
        "（降幅88.7%）且最优纯目标相对变化约0.01%；"
    )
    for old in abstract_old:
        _replace_in_doc(doc, old, abstract_new)


def _sync_table1(doc: Document) -> None:
    """按 patent_experiment_tables 刷新说明书表1（3 行场景）。"""
    headers, rows = paired_table_rows()
    if not headers or not rows:
        return
    for table in doc.tables:
        if not table.rows:
            continue
        if table.rows[0].cells[0].text.strip() != "场景":
            continue
        for i, h in enumerate(headers):
            if i < len(table.rows[0].cells):
                _write_cell(table.rows[0].cells[i], h)
        while len(table.rows) > len(rows) + 1:
            table._tbl.remove(table.rows[-1]._tr)
        while len(table.rows) < len(rows) + 1:
            table.add_row()
        for ri, row_data in enumerate(rows):
            for ci, val in enumerate(row_data):
                if ci < len(table.rows[ri + 1].cells):
                    _write_cell(table.rows[ri + 1].cells[ci], val)
        return


def patch_existing_docx() -> Document:
    if not NEW_PATH.exists():
        raise FileNotFoundError(f"找不到已编辑的专利文件：{NEW_PATH}")
    doc = Document(str(NEW_PATH))
    apply_shared_text_patches(doc)
    return doc


def main() -> None:
    import argparse

    parser = argparse.ArgumentParser(description="生成或修订专利 docx 正文")
    parser.add_argument(
        "--patch-only",
        action="store_true",
        help="仅修订现有 docx 文字（保留已嵌入附图），不从头重建",
    )
    args = parser.parse_args()

    if args.patch_only:
        if not NEW_PATH.exists():
            raise SystemExit(f"[ERR] --patch-only 需要已存在：{NEW_PATH}")
        doc = patch_existing_docx()
        doc.save(NEW_PATH)
        print(f"[OK] 已修订附图相关正文：{NEW_PATH}")
        return

    if NEW_PATH.exists():
        doc = patch_existing_docx()
        doc.save(NEW_PATH)
        print(f"[OK] 检测到已有 docx，已修订文字并保留附图：{NEW_PATH}")
        return

    new_doc = build_new_docx()
    new_doc.save(NEW_PATH)
    print(f"[OK] 已基于底稿生成新专利：{NEW_PATH}")
    print(f"[INFO] 底稿：{ORIGINAL_1110}")


if __name__ == "__main__":
    main()
